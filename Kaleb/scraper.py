import requests
from bs4 import BeautifulSoup
import trafilatura
import logging
import json
import os
from typing import List, Set
from urllib.parse import urljoin, urlparse
import io
from pypdf import PdfReader
import pandas as pd
from docx import Document

logger = logging.getLogger(__name__)

class ContentScraper:
    """
    A configured web scraper designed to crawl utility websites for energy efficiency incentives.
    
    Design Philosophy:
    - Targeted Crawling: Instead of a broad crawl, it starts from a specific URL and only follows links 
      that match specific keywords or patterns (e.g. "program", "incentive").
    - Depth Control: strictly limited by `max_pages` to prevent getting lost in large corporate sites.
    - Content Extraction: Uses `trafilatura` for main text and `pypdf` for PDF documents, as incentive 
      details are often hidden in PDF downloads.
    """
    def __init__(self, max_pages: int = 5):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        self.max_pages = max_pages
        # Load keywords from config
        self.keywords = self._load_keywords()

    def _load_keywords(self) -> List[str]:
        try:
            config_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'config', 'keywords.json')
            if os.path.exists(config_path):
                with open(config_path, 'r') as f:
                    data = json.load(f)
                    return data.get('search_keywords', [])
            return ['incentive', 'rebate', 'efficiency', 'solar', 'renewable'] # fallback
        except Exception as e:
            logger.error(f"Error loading keywords: {e}")
            return ['incentive', 'rebate']

    def fetch_content(self, base_url: str) -> str:
        """
        Crawls the website starting from base_url, looking for incentive-related pages.
        Returns combined text content from relevant pages.
        
        Args:
            base_url: The starting URL. This may be a specific page or a general landing page.
            
        Returns:
            str: A large string containing the concatenated text of all visited pages, 
                 separated by "--- SOURCE: [url] ---" headers.
        """
        if not base_url:
            return ""

        visited: Set[str] = set()
        to_visit: List[str] = [base_url]
        collected_text = []
        
        # Use loaded keywords
        keywords = self.keywords
        
        start_netloc = None

        count = 0
        while to_visit and count < self.max_pages:
            url = to_visit.pop(0)
            if url in visited:
                continue
            
            try:
                logger.info(f"Scraping: {url}")
                response = self.session.get(url, timeout=10)
                response.raise_for_status()
                visited.add(url)
                count += 1
                
                content_type = response.headers.get('Content-Type', '').lower()

                # Handle PDF
                if 'application/pdf' in content_type or url.lower().endswith('.pdf'):
                    pdf_text = self._extract_pdf_text(response.content)
                    if pdf_text:
                        collected_text.append(f"--- SOURCE (PDF): {url} ---\n{pdf_text}\n")
                    continue # Stop processing this URL (no links to extract from PDF usually)

                # Handle HTML
                text = trafilatura.extract(response.text)
                if text:
                    collected_text.append(f"--- SOURCE: {url} ---\n{text}\n")

                # If this is the first successful page, determine the primary domain.
                # REASONING: We cannot rely on the input `base_url` for domain warnings
                # because many inputs are shortlinks (bit.ly) or redirects. 
                # We need the *final* destination URL to know what domain we are actually on.
                current_netloc = urlparse(response.url).netloc
                if start_netloc is None:
                    start_netloc = current_netloc
                
                # If we are on the homepage or an index page, look for sub-pages and PDFs
                # Limiting depth implicitly by only adding filtered links
                soup = BeautifulSoup(response.text, 'html.parser')
                for link in soup.find_all('a', href=True):
                    href = link['href']
                    # Use the actual response URL to resolve relative links, not the initial base_url.
                    # REASONING: If we were redirected from `bit.ly/xyz` to `site.com/page`, 
                    # a link `<a href="subpage">` must be resolved against `site.com/page`, 
                    # not `bit.ly/xyz`.
                    full_url = urljoin(response.url, href)
                    
                    # check if it matches the primary domain we are scraping
                    link_domain = urlparse(full_url).netloc
                    
                    if link_domain == start_netloc:
                        # Skip typical non-content files (images, zips) but ALLOW PDFs
                        if any(ext in full_url.lower() for ext in ['.jpg', '.png', '.zip', '.exe']):
                            continue

                        link_text = link.get_text().lower()
                        href_lower = href.lower()
                        
                        is_relevant = any(k in link_text for k in keywords) or any(k in href_lower for k in keywords)
                        
                        if is_relevant:
                            if full_url not in visited and full_url not in to_visit:
                                to_visit.append(full_url)
                        # Heuristic: If the link text contains "read more", "details", "click here", assume it might be relevant if near relevant content context, 
                        # but for now let's just stick to the expanded keywords and trust them.
                        # What we WILL do is check if the URL itself looks very "program-like" even if keywords missed it (which they shouldn't now).
                        elif "/program" in full_url.lower() or "/incentive" in full_url.lower():
                             if full_url not in visited and full_url not in to_visit:
                                to_visit.append(full_url)
            
            except Exception as e:
                logger.error(f"Error scraping {url}: {e}")
        
        return "\n".join(collected_text)

    def _extract_pdf_text(self, content_bytes: bytes) -> str:
        try:
            with io.BytesIO(content_bytes) as f:
                reader = PdfReader(f)
                text = []
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text.append(extracted)
                return "\n".join(text)
        except Exception as e:
            logger.error(f"Failed to extract PDF text: {e}")
            return ""

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    scraper = ContentScraper()
    #print(scraper.fetch_content("hhttp://comptroller.marylandtaxes.com/Public_Services/Agency_Information/Office_of_the_Comptroller/Comptroller_Initiatives/Shop_Maryland_Tax-free_Week/"))

    scraped_content = []
    relevant_urls = pd.read_excel('Relevant URLs.xlsx', header = None)
    scraped_content = []
    # Iterate through URLs and scrape content form the first 10 URLs
    for url in relevant_urls[0][:10]: 
        print(f"Fetching content for: {url}")
        try:
            content = scraper.fetch_content(url)
            scraped_content.append(content) # save it
        except Exception as e:
            print(f"Error fetching content for {url}: {e}")
    #Save to Word document
    print("Writing to document...")
    doc = Document()
    doc.add_heading('Scraped Content Report', 0)
    for url, content in zip(relevant_urls[0], scraped_content):
        doc.add_heading(f'URL: {url}', level=1)
        doc.add_paragraph(content)

    doc.save('Scraped_Results.docx')
    print("Done!")
